from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).parent / "SignGuyAI_Rebuild_Combined_Plan_and_Conflict_Register.docx"

NAVY = "17324D"
TEAL = "128C8C"
LIGHT_TEAL = "E7F4F3"
LIGHT_BLUE = "EAF0F6"
LIGHT_GRAY = "F2F4F7"
GOLD = "FFF2CC"
RED = "FCE8E6"
GREEN = "E6F4EA"
WHITE = "FFFFFF"
INK = "1F2933"
MUTED = "5B6770"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, bold=False, color=INK, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_borders(table, color="CBD5E1", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, label, text, fill=LIGHT_TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_simple_table(doc, headers, rows, widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        set_cell_text(hdr.cells[i], value, bold=True, color=WHITE, size=font_size)
        shade(hdr.cells[i], NAVY)
        if widths:
            hdr.cells[i].width = Inches(widths[i])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=font_size)
            shade(cells[i], WHITE if row_index % 2 == 0 else LIGHT_GRAY)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.16)

    for level, size, color in ((1, 16, NAVY), (2, 13, TEAL), (3, 11, NAVY)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("SignGuyAI Rebuild Combined Plan | June 13, 2026")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(55)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("SignGuyAI Clean Rebuild")
    r.font.name = "Aptos Display"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Combined Master Plan and Conflict Register")
    r.font.name = "Aptos Display"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)

    add_callout(
        doc,
        "Controlling interpretation",
        "Any explicit owner/user-added decision is definite. Where documents still disagree, this report preserves the definite outcome, applies the latest specific technical lock, and flags only the choices that still require owner confirmation.",
        LIGHT_TEAL,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Prepared from the ten supplied rebuild documents")
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    add_body(doc, "Prepared June 13, 2026. This is a planning synthesis, not an implementation audit of the current repository.")

    add_body(
        doc,
        "Important duplicate: “chat webstores.pdf” and “signguy_ai_rebuild_addon_strategy.pdf” are byte-for-byte identical. They are treated as one source, not two independent votes.",
    )
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. Executive Combined Decision", 1)
    add_body(
        doc,
        "Rebuild SignGuyAI as a clean modular monolith in a new repository. The current app is the product-behavior, visual-reference, and migration source; its structural debt is not the architecture template.",
    )
    add_bullets(
        doc,
        [
            "Preserve owner-approved behavior, fields, defaults, calculations, workflows, labels, and design behavior unless the owner explicitly changes them.",
            "Establish the final app surface early, but distinguish visible structure from fully activated depth.",
            "Build a useful Base App plus cleanly gated paid add-ons. Webstores is the first full expansion add-on; Wrap Command Center is second.",
            "Use shared platform systems for customers, orders, documents, forms, portals, notifications, permissions, payments, tasks, reporting, and AI usage. Add-ons add specialized workflows, not duplicate foundations.",
            "Apply the Rebuilt Data Layer Blueprint final locks to authoritative collections.",
            "Use deterministic services for pricing, inventory, payroll, taxes, accounting, payments, refunds, permissions, and approvals. AI may assist but may not silently mutate these areas.",
            "Use feature flags, acceptance tests, migration reconciliation, rollback plans, and staged tenant rollout for every major capability.",
        ],
    )

    add_heading(doc, "2. Authority and Conflict Rules", 1)
    add_simple_table(
        doc,
        ["Priority", "Source / Rule", "How It Controls the Combined Plan"],
        [
            ("1", "Explicit owner/user-added decisions", "Definite. They override generic simplification, old sequencing, and draft examples."),
            ("2", "Latest specific final locks", "The v2 master-plan updates and Data Layer Blueprint final locks control where they are more specific."),
            ("3", "Emergent “Know from Day 1” guidance", "Controls implementation guidance and protected approved behavior unless overridden by security, data integrity, or owner decisions."),
            ("4", "Accepted conflict resolutions", "Controls the original 15 known contradictions unless a newer final lock supersedes one."),
            ("5", "Older roadmap and chat-plan text", "Useful for scope and sequencing only when it does not conflict with higher-priority decisions."),
        ],
        widths=[0.55, 2.0, 3.9],
    )
    add_callout(
        doc,
        "Non-negotiable exception",
        "Tenant isolation, security, payment safety, deterministic financial/inventory/payroll math, auditability, and explicit owner-approved behavior override generic implementation suggestions.",
        GOLD,
    )

    add_heading(doc, "3. Product Structure and Packaging", 1)
    add_heading(doc, "3.1 Base App", 2)
    add_body(doc, "The Base App must run a sign shop on its own. It includes:")
    add_bullets(
        doc,
        [
            "Home / Global Command Center, shared layout, contextual ribbon system, global search, global create, notifications, help, profile, settings, permissions, and feature gating.",
            "Customers/CRM, orders, quotes, invoices, manual payments, basic reports, basic pricing, basic production board, basic team features, and productivity basics.",
            "Basic customer portal, document sharing, questionnaires/forms, email, and shared payment/portal/notification foundations.",
            "Visible module states for final product areas, using enabled, disabled, setup_required, coming_soon, or unavailable instead of fake zeroes.",
        ],
    )
    add_heading(doc, "3.2 Gated Add-Ons", 2)
    add_simple_table(
        doc,
        ["Order", "Add-On", "Combined Scope Rule"],
        [
            ("1", "Webstores", "First complete production-ready expansion module; preserves approved current behavior and proves shared portals, products, orders, payments, forms, documents, and notifications."),
            ("2", "Wrap Command Center", "Full wrap workflow from intake and vehicle records through quote, proof, production, install, and payment."),
            ("3", "Smart Pricing Pro", "Advanced pricing analysis, detailed modes, live updates, market research, margin warnings, and AI recommendations; Base App retains usable basic pricing."),
            ("4+", "AI Tools Pro, SMS/MMS, Production Timing Pro, Training Pro", "Advanced or usage-cost capabilities layered on shared foundations and gated independently."),
        ],
        widths=[0.5, 1.6, 4.35],
    )
    add_body(
        doc,
        "Feature gating controls entitlement and activation. It must not create duplicate customer, order, payment, portal, document, message, product, task, or reporting models.",
    )

    add_heading(doc, "4. Final App Surface", 1)
    add_callout(
        doc,
        "Recommended controlling shell",
        "Use the v2 accepted dashboard baseline: a dark compact workspace rail ordered Home, Operations, Business, Productivity, AI Hub, Settings; global top bar; contextual module navigation; and one reusable contextual ribbon system.",
    )
    add_simple_table(
        doc,
        ["Surface", "Primary Responsibility"],
        [
            ("Home / Global Command Center", "Cross-workspace Action Required, KPIs, snapshots, onboarding, and reserved future cards. It is not the Business dashboard."),
            ("Operations", "Customers, quotes, orders, order items/job-ticket views, production, proofs, artwork, installs, documents, and operational Webstores access."),
            ("Business", "Invoices, payments, AR, sales, expenses, taxes, deductions, payroll, purchasing/inventory alerts, reports, and financial controls."),
            ("Productivity", "Tasks, Kanban, reminders, internal messages, announcements, calendar, schedules, and checklists."),
            ("AI Hub", "Help, onboarding, documentation, prompt library, feedback, roadmap, release notes, and later guarded AI tools."),
            ("Settings", "Tenant configuration, permissions, pricing foundation, integrations, feature/add-on controls, and domain settings."),
        ],
        widths=[1.8, 4.65],
    )
    add_body(
        doc,
        "Every workspace dashboard uses the same shell and card language but has contextual cards and actions. Unavailable capabilities remain visible only where that helps establish the final structure, and must honestly show their state.",
    )

    add_heading(doc, "5. Architecture and Shared Systems", 1)
    add_bullets(
        doc,
        [
            "One deployable frontend and backend initially, with enforced domain boundaries.",
            "Routes handle authentication, validation, request/response, and error mapping. Domain services own business rules. Repositories own persistence.",
            "No giant server.py, giant page components, giant AppContext, direct route DB access as the normal pattern, process-local scheduling, unmanaged indexes, or base64-heavy file storage.",
            "Frontend routes use lazy loading and code splitting. Dashboards use digest-style endpoints, indexed aggregation, pagination, and reusable widgets.",
            "Object storage uses file records, signed URLs, streaming, and non-blocking paths where appropriate.",
            "Background jobs are idempotent, observable, retryable, and safe across multiple workers.",
            "Public/portal surfaces use separate route and API boundaries with allowlisted data. Internal costs and restricted data never leak.",
        ],
    )
    add_heading(doc, "5.1 Canonical Shared Systems", 2)
    add_body(
        doc,
        "Use one shared implementation for each cross-cutting concept: customers, contacts, tasks, notes/comments, messages/conversations, calendar events, notifications, files/assets, questionnaires/forms, activity timelines, audit events, permissions, feature flags, portal framework, payment foundation, reporting foundation, and AI usage tracking.",
    )
    add_heading(doc, "5.2 Protected Behavior", 2)
    add_body(doc, "Before rebuilding each protected area, create and approve a behavior specification with acceptance tests:")
    add_bullets(
        doc,
        [
            "Pricing Foundation, product/item categories, category quizzes/progressive forms, defaults, and calculation assumptions.",
            "Onboarding setup, Office-style ribbon behavior, dashboard layouts, global create, and global search.",
            "Webstore setup, store types, owner onboarding, catalog behavior, storefront, order sync, payouts, and portals.",
            "Team schedule, timeclock, payroll, inventory/purchasing, Wrap Command Center, document library/questionnaires, portals, and AI guardrails.",
        ],
    )

    add_heading(doc, "6. Controlling Data-Layer Decisions", 1)
    add_callout(
        doc,
        "Final technical lock",
        "The Rebuilt Data Layer Blueprint supersedes older draft schema examples where it states a decision or final lock.",
        LIGHT_BLUE,
    )
    add_simple_table(
        doc,
        ["Area", "Controlling Decision"],
        [
            ("Tenant scope", "Every tenant-owned authoritative document contains tenant_id; all tenant-scoped unique indexes begin with tenant_id."),
            ("IDs", "Application-generated UUIDv7 strings; MongoDB _id remains internal and is never exposed."),
            ("Dates", "Persist native timezone-aware UTC BSON datetimes, not ISO date strings, in authoritative collections."),
            ("Money", "Persist integer minor units/cents; rates and percentages use explicit rate representations."),
            ("Work model", "orders and order_items are the only authoritative sellable-work aggregates. “Job ticket” may remain a UI label for a production order item; no separate authoritative JobTicket aggregate."),
            ("Billing lines", "Fees, deposits, discounts, delivery, permits, and other non-produced charges belong in Billing-owned quote/invoice lines rather than creating production work records."),
            ("Subscriptions", "Versioned plan_catalog + subscriptions + entitlement_grants + usage_counters. Founder status is a commercial grant/price lock, not a feature-bypass flag."),
            ("Pricing", "Versioned pricing_profiles + immutable pricing_snapshots. Historical snapshots preserve sold/quoted truth."),
            ("References", "No mutable cross-domain aggregates embedded in another domain; references live on the dependent side; no unbounded arrays or authoritative bidirectional lists."),
            ("Indexes", "Manifest-driven, named, validated, migration-controlled indexes; required unique/security indexes fail readiness."),
        ],
        widths=[1.35, 5.1],
    )

    add_heading(doc, "7. Webstores Combined Plan", 1)
    add_body(
        doc,
        "Webstores are protected approved behavior and the first full expansion module. They are not a shell-only future idea and are not automatically deferred to Release 6.",
    )
    add_numbered(
        doc,
        [
            "Build the app shell and shared core contracts first.",
            "Create and approve the protected Webstore Behavior Spec and acceptance tests.",
            "Preserve active approved setup, store types, owner invite/onboarding, questionnaires/uploads, products/catalog, public storefront, cart, order sync, dashboard/ribbon placement, and portal behavior.",
            "Implement Webstores in vertical slices. Feature-flag only unsafe dependencies such as unproven payouts, advanced inventory reservations, supplier integrations, or incomplete external services.",
            "At checkout, create a WebstoreOrder and bridge purchased items into canonical orders/order_items and Billing/Production records as appropriate.",
            "Use the shared portal, document, questionnaire, email, notification, payment, reporting, and permission foundations.",
        ],
    )
    add_heading(doc, "7.1 Webstore Product Boundary", 2)
    add_body(
        doc,
        "Webstore products remain ecommerce catalog records, not InventoryItems and not authoritative shop-order records. They may reference shared product templates, pricing profiles, materials, and production templates. Purchases map into canonical order_items and billing lines.",
    )
    add_heading(doc, "7.2 Payment Safety", 2)
    add_body(
        doc,
        "Stripe Connect onboarding, webhook verification, idempotency, tenant isolation, reconciliation, refunds, disputes, disconnect recovery, and failure recovery are required before live creator/store payout behavior is enabled.",
    )

    add_heading(doc, "8. Combined Execution Sequence", 1)
    add_simple_table(
        doc,
        ["Phase", "Primary Outcome", "Required Scope"],
        [
            ("0", "Decisions and foundation", "Protected specs; final route/navigation/permission map; add-on entitlement map; canonical data models; migration mappings; repo structure; CI; indexes; auth; audit; observability; fixtures."),
            ("1", "Base shell and pilot core", "Official Home dashboard baseline; final workspace surfaces; settings; permissions; search/create; customers; basic pricing; quotes; orders/order_items; production views; documents; invoices/manual payments; basic portal/productivity."),
            ("1W", "Webstores preservation track", "Runs in parallel where practical: preserve approved setup, catalog, owner flow, questionnaires, storefront, cart, and canonical order bridge; gate unsafe payment dependencies."),
            ("2", "Webstores full expansion", "First complete gated add-on with production-ready shared systems, safe checkout, owner portal, reporting, and approved store modes."),
            ("3", "Wrap Command Center", "Second complete gated add-on using shared intake, files, pricing, approvals, scheduling, production, and portal systems."),
            ("4", "Workforce, inventory, and financial depth", "Deepen payroll, team, inventory ledger, purchasing, financial controls, reports, and communications according to dependency readiness."),
            ("5+", "Advanced paid modules", "Smart Pricing Pro, AI Tools Pro, SMS/MMS, Production Timing Pro, Training Pro, and later ecosystem capabilities."),
        ],
        widths=[0.55, 1.75, 4.15],
    )
    add_body(
        doc,
        "This sequence replaces contradictory release numbers with one rule: final structure and contracts early; Base App usable first; Webstores first full add-on; then other add-ons and deeper domains by dependency and risk.",
    )

    add_heading(doc, "9. Conflict Register", 1)
    add_body(
        doc,
        "Status meanings: Resolved means the supplied documents provide a controlling answer under the authority rules. Owner confirmation means two definite-looking product choices still conflict and cannot be safely inferred.",
    )
    conflicts = [
        ("1", "Webstores timing", "Late Release 6 activation vs early protected active scope.", "Resolved", "Webstores is early and is the first full expansion add-on. Release 6 language applies only to later ecommerce depth."),
        ("2", "Final surface vs one-module focus", "All final areas visible early vs Webstores built first and deeply.", "Resolved", "Build all final surfaces/contracts early; make Webstores the first complete deep add-on."),
        ("3", "Workspace navigation labels", "Four workspaces: Operations, Business, Productivity, AI Hub vs top-level Home, Operations, Business, Team, Webstores.", "Owner confirmation", "Recommended: use v2 rail Home, Operations, Business, Productivity, AI Hub, Settings. Keep Team inside Business/Productivity and Webstores as a gated module."),
        ("4", "Permanent left navigation", "Chat plan prohibits permanent whole-app left sidebar; v2 accepts a dark left workspace rail.", "Resolved", "The later v2 dashboard baseline controls: use a compact permanent workspace rail, not a large permanent module sidebar."),
        ("5", "Webstores navigation placement", "Webstores as top-level navigation vs module/add-on inside the workspace structure.", "Owner confirmation", "Recommended: entitlement-aware Webstores entry inside Operations plus optional pinned rail shortcut when enabled; do not create a fifth competing workspace without approval."),
        ("6", "Team navigation placement", "Team as top-level area vs Productivity and Business ownership.", "Owner confirmation", "Recommended: tasks/messages/calendar in Productivity; employees/timeclock/payroll in Business; shared Team landing can deep-link to both."),
        ("7", "Base App vs paid add-ons", "Protected features must remain present, while Webstores/Wrap/advanced tools are paid add-ons.", "Resolved", "Preserve target behavior and contracts; gate entitlement/activation. Base App remains useful and add-ons do not duplicate core systems."),
        ("8", "JobTicket data model", "Separate JobTicket aggregate vs Data Blueprint says job ticket becomes order_item/UI terminology.", "Resolved", "Data Blueprint final lock controls: order_items are authoritative; retain Job Ticket as a UI/workflow label where useful."),
        ("9", "Produced items vs fees", "Every line item/job ticket vs production-only work.", "Resolved", "Production order_items represent produced deliverables. Billing-owned lines represent fees, discounts, deposits, delivery, permits, and similar charges."),
        ("10", "Timestamp storage", "ISO strings in earlier schemas vs native UTC BSON datetimes.", "Resolved", "Use native timezone-aware UTC BSON datetimes in authoritative collections; serialize at API boundaries."),
        ("11", "ID format", "Generic UUID/UUID4 examples vs UUIDv7 final rule.", "Resolved", "Use application-generated UUIDv7 strings; temporary UUID4 example is implementation scaffolding only."),
        ("12", "Tier/founder gating", "Persist tier features and is_founder bypass vs normalized plans/entitlements/grants.", "Resolved", "Use plan catalog, subscriptions, entitlement grants, and usage counters. Founder status is a price lock/grant, never a global bypass."),
        ("13", "Settings storage", "One tenant_settings document forever vs domain-owned settings.", "Resolved", "Expose one merged settings read model; store settings by domain when scale, permissions, or ownership require it."),
        ("14", "Product catalog boundary", "One reusable product/catalog system vs separate WebstoreProduct and InventoryItem.", "Resolved", "Share catalog capabilities and references, but keep bounded authoritative records. WebstoreProduct is ecommerce; InventoryItem is stock; order_item is sold work."),
        ("15", "Stripe Connect charge model", "Owners receive funds directly when appropriate vs platform charge then distribute minus fee.", "Owner confirmation", "Choose and document the Connect charge type per store mode before implementation; model fees, refunds, disputes, tax, reconciliation, and liability around that choice."),
        ("16", "Webstore type taxonomy", "Fundraiser/B2B/creator vs fundraiser/promotional/B2B/general/event.", "Owner confirmation", "Recommended: define canonical store_type values and treat promotional/event/general as configurable campaign templates unless they truly have different rules."),
        ("17", "AI timing", "AI framework/identity early vs active or mutating AI later.", "Resolved", "AI Hub, usage tracking, prompt/help surfaces, and contracts early; mutating actions only after deterministic services and review guardrails."),
        ("18", "Portal timing", "Separate boundaries and basic portal early vs full portals later.", "Resolved", "Build shared portal framework, security boundaries, and Base customer portal early; deepen specialized portals by dependency/add-on."),
        ("19", "SMS timing", "SMS/MMS foundation in shared core vs paid messaging add-on later.", "Resolved", "Build provider abstraction, permissions, data model, and notification contracts early; gate live sending, two-way messaging, and usage billing."),
        ("20", "Inventory timing", "Protected/current integration vs delayed full inventory release.", "Resolved", "Define models/contracts/material links early; activate full ledger and purchasing after core order/production stability unless migration requires earlier behavior."),
        ("21", "Dashboard ownership", "Mixed global dashboard shown as Business vs Home / Command Center.", "Resolved", "The accepted screenshot becomes Home-active. Business receives contextual financial/workforce/purchasing cards."),
        ("22", "Inactive dashboard cards", "Return zeroes, hide cards, or reserve stable cards.", "Resolved", "Reserve useful final-surface cards and return explicit module state with null counts; never imply inactive capability with zero."),
        ("23", "Current behavior preservation vs quarantine", "Preserve features vs quarantine unreliable external/unfinished flows.", "Resolved", "Preserve approved target behavior/specification; do not activate unreliable implementation until security and acceptance tests pass."),
        ("24", "Release numbering", "Older Releases 0-8 vs App Shell/Shared Core/Add-On phases.", "Resolved", "Use the combined execution sequence in this document. Old release numbers become dependency context, not controlling milestones."),
        ("25", "Pricing scope", "Pricing Foundation/core calculator vs Smart Pricing Pro paid add-on.", "Resolved", "Base includes usable deterministic pricing foundation and calculator; Pro contains advanced analysis, live intelligence, market research, and AI recommendations."),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    headers = ["#", "Conflict", "Contradiction", "Status", "Combined Outcome / Solution"]
    widths = [0.3, 1.1, 1.65, 0.85, 2.55]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=WHITE, size=8)
        shade(table.rows[0].cells[i], NAVY)
        table.rows[0].cells[i].width = Inches(widths[i])
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(conflicts):
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=7.6)
            if i == 3:
                shade(cells[i], GOLD if value == "Owner confirmation" else GREEN)
            else:
                shade(cells[i], WHITE if idx % 2 == 0 else LIGHT_GRAY)
            cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    add_heading(doc, "10. Owner Confirmations Still Required", 1)
    add_callout(
        doc,
        "Only five decisions remain genuinely open",
        "Everything else in the conflict register has a controlling resolution from the supplied documents.",
        GOLD,
    )
    add_numbered(
        doc,
        [
            "Approve the final workspace rail and module placement: recommended Home, Operations, Business, Productivity, AI Hub, Settings.",
            "Confirm whether Webstores gets an always-visible top-level rail item, an entitlement-aware pinned item, or remains an Operations module.",
            "Confirm whether Team gets a shared landing page and where it appears in navigation.",
            "Choose the Stripe Connect charge model for each store mode and document liability/reconciliation rules.",
            "Approve the canonical webstore type taxonomy and which differences are true business rules versus templates.",
        ],
    )

    add_heading(doc, "11. Definition of Done", 1)
    add_bullets(
        doc,
        [
            "Protected behavior specifications and acceptance tests exist before replacement implementation.",
            "Tenant isolation, RBAC, portal allowlists, audit events, and permission-filtered navigation/actions are tested for every domain.",
            "Money, inventory, payroll, tax, payment, refund, and approval logic is deterministic and server-authoritative.",
            "Every page has loading, empty, error, success, confirmation, setup/disabled, and permission-denied states as applicable.",
            "No N+1 or repeated large in-memory collection reads in dashboards, productivity, search, or reports.",
            "Migrations include deterministic ID maps, count/financial reconciliation, rollback windows, and read-only legacy retention until verified.",
            "Frontend unit/component and browser workflow tests exist; backend tests are deterministic and do not rely only on live URLs.",
            "External integrations are enabled only after webhook verification, idempotency, disconnect/recovery behavior, and observability are proven.",
        ],
    )

    add_heading(doc, "12. Source Document Treatment", 1)
    add_simple_table(
        doc,
        ["Source", "Treatment in This Combined Plan"],
        [
            ("CONFLICT DECISIONS.pdf", "Original 15-conflict worksheet and accepted synthesis rules."),
            ("MASTER PLAN WITH CONFLICT RESOLUTION.pdf", "Accepted combined product/architecture/release baseline."),
            ("REBUILD MASTER PLAN.pdf", "Later v2 update; controls dashboard baseline, audit guardrails, and source priority."),
            ("Clean Rebuild and Release Roadmap.pdf", "Broad product/release/detail source; contradictory old release/schema examples are superseded."),
            ("chat plan.pdf", "Final-surface, protected behavior, architecture, and first-version requirements."),
            ("chat plan with protected areasd.pdf", "Protected-behavior version of the chat plan; treated as reinforcing the protected-spec rule."),
            ("signguy_ai_rebuild_strategy_app_shell_webstores.pdf", "Definite app-shell-first and Webstores-first-expansion strategy; conflicting nav labels flagged."),
            ("signguy_ai_rebuild_addon_strategy.pdf", "Definite Base App + clean paid add-on packaging and order."),
            ("chat webstores.pdf", "Exact duplicate of signguy_ai_rebuild_addon_strategy.pdf; counted once."),
            ("codex Rebuilt Data Layer Blueprint.pdf", "Controlling detailed data decisions and final locks."),
        ],
        widths=[2.65, 3.8],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
